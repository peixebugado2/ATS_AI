import os
import json
import re
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Tuple

import streamlit as st
from dotenv import load_dotenv
from google import genai
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

load_dotenv()

APP_TITLE = "CareerOps Studio"
APP_SUBTITLE = "Executive Resume Optimization Platform"
DEFAULT_MODEL = "gemini-2.5-flash-lite"
ROOT = Path(__file__).parent
TEMPLATE_DIR = ROOT / "templates"

HISTORY_FILE = ROOT / "careerops_history.json"
MAX_HISTORY_ITEMS = 30

MODEL_OPTIONS = [
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash",
    "gemini-2.0-flash-lite",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-flash-8b",
]

CONTEXT_PROFILES = {
    "Fast": {"instructions": 8000, "master": 26000, "manager": 8000, "director": 8000},
    "Balanced": {"instructions": 14000, "master": 48000, "manager": 12000, "director": 12000},
    "Deep": {"instructions": 24000, "master": 80000, "manager": 18000, "director": 18000},
}

ROLE_NAMES = [
    "Sr. EMEA Transport Manager",
    "Director Logistics",
    "Trade Compliance Manager EMEA",
    "Head of Logistics EMEA",
    "Sr. EMEA Logistics & Spare Parts Inventory Manager",
]

EXPECTED_COUNTS = {
    "Director CV": {
        "Sr. EMEA Transport Manager": 7,
        "Director Logistics": 8,
        "Trade Compliance Manager EMEA": 0,
        "Head of Logistics EMEA": 7,
        "Sr. EMEA Logistics & Spare Parts Inventory Manager": 5,
    },
    "Manager CV": {
        "Sr. EMEA Transport Manager": 7,
        "Director Logistics": 0,
        "Trade Compliance Manager EMEA": 8,
        "Head of Logistics EMEA": 7,
        "Sr. EMEA Logistics & Spare Parts Inventory Manager": 5,
    },
}

REQUIRED_CV_SECTIONS = [
    "Name",
    "Contact Details",
    "Additional ATS Skills",
    "Professional Summary",
    "Professional Experience",
    "Education",
    "Systems / Tools",
    "ATS Analysis",
]

MANDATORY_CV_TYPE_RULES = """
CV TYPE SELECTION RULES - NON NEGOTIABLE

Director CV:
Use when the Job Description mentions Director, Head of, Senior Leadership, Strategic Leadership, Global Leadership, Regional Leadership, executive ownership or director-level accountability.
Mandatory sections: Sr. EMEA Transport Manager; Director Logistics; Head of Logistics EMEA; Sr. EMEA Logistics & Spare Parts Inventory Manager.
Director CV must NOT include Trade Compliance Manager EMEA.

Manager CV:
Use when the Job Description mentions Manager, Operations Manager, Supply Chain Manager, Logistics Manager, Warehouse Manager or operational management.
Mandatory sections: Sr. EMEA Transport Manager; Trade Compliance Manager EMEA; Head of Logistics EMEA; Sr. EMEA Logistics & Spare Parts Inventory Manager.
Manager CV must include Trade Compliance Manager EMEA instead of Director Logistics.
"""

MANDATORY_BULLET_COUNTS = """
MANDATORY ACHIEVEMENT COUNTS - NON NEGOTIABLE

Director CV:
- Sr. EMEA Transport Manager: exactly 7 achievement lines
- Director Logistics: exactly 8 achievement lines
- Trade Compliance Manager EMEA: must not appear
- Head of Logistics EMEA: exactly 7 achievement lines
- Sr. EMEA Logistics & Spare Parts Inventory Manager: exactly 5 achievement lines

Manager CV:
- Sr. EMEA Transport Manager: exactly 7 achievement lines
- Director Logistics: must not appear
- Trade Compliance Manager EMEA: exactly 8 achievement lines, not 7
- Head of Logistics EMEA: exactly 7 achievement lines, not 6
- Sr. EMEA Logistics & Spare Parts Inventory Manager: exactly 5 achievement lines

Do not add extra roles. Do not omit required roles. Do not create fewer or more achievement lines.
"""

STRICT_FORMAT_RULES = """
STRICT COPY-PASTE FORMAT RULES - MACHINE VALIDATED

The output must be compact, ATS optimized and easy to copy and paste into the existing CV template.

Mandatory section order inside FINAL TAILORED CV:
1. Name
2. Contact Details
3. Additional ATS Skills
4. Professional Summary
5. Professional Experience
6. Education
7. Systems / Tools
8. ATS Analysis
9. Cover Letter only if requested

Additional ATS Skills formatting:
- Exactly 10 skills.
- One skill per line.
- No bullets.
- No numbering.
- No blank lines between skills.
- No comma-separated skills.
- Compact vertical list only.

Professional Experience formatting:
- Each role heading must be on its own line.
- Each achievement must be on its own separate new line.
- DO NOT write "ACHIEVEMENT:".
- DO NOT write "Achievement:".
- DO NOT write "Bullet:".
- DO NOT write "Result:".
- DO NOT write "Success:".
- DO NOT write "KPI:".
- DO NOT use bullet symbols such as -, *, •, or numbered bullets for achievements.
- Do not place multiple achievements on the same paragraph.
- One blank line between roles only.
- No blank lines between achievements inside the same role.
- Each achievement must occupy exactly one line.
- Achievement text must be plain text only, ready to paste into the CV.

Professional experience role headings must match exactly:
Sr. EMEA Transport Manager
Director Logistics
Trade Compliance Manager EMEA
Head of Logistics EMEA
Sr. EMEA Logistics & Spare Parts Inventory Manager

Every achievement line must be 170 to 190 characters including spaces.
"""

COVER_LETTER_RULES = """
COVER LETTER RULES - MACHINE VALIDATED

The cover letter must be generated from the Job Description and the selected verified achievements.
Extract the company name, exact job title, location, seniority, department, responsibilities, systems, hard skills, soft skills and ATS keywords from the Job Description whenever available.
Never use placeholders such as [Company Name], [Hiring Manager Name], [Job Title] or similar text.
If company name is available, address it to: Dear Hiring Team at [Company Name],
If company name is not available, address it to: Dear Hiring Manager,
The cover letter body must be 1,795 to 1,805 characters including spaces. Target exactly 1,800 characters.
The Cover Letter must be written in the same language as the Job Description unless the user explicitly selected another output language.
Do not show the character count in the final output.
Do not invent facts, achievements, metrics, certifications, systems or company information.
"""

ATS_RULES = """
ATS VALIDATION RULES - NO LIVE SITE CLAIMS

The ATS validation is an internal simulation based only on the Job Description, generated CV and source files.
It is not connected to LinkedIn, Jobscan, SkillSyncer, Resume Worded, Rezi or any external ATS website.
Never claim that the score came directly from live websites.
Use wording like: simulated ATS-style score, internal ATS simulation, LinkedIn-style, Jobscan-style, SkillSyncer-style.

Skill extraction requirements:
- Extract hard skills directly from the Job Description.
- Extract soft skills directly from the Job Description.
- Extract tools, systems, certifications, industry terms, leadership requirements and operational keywords directly from the Job Description.
- Match extracted skills against the tailored CV.
- Additional ATS Skills must contain exactly 10 skills and must be aligned with the Job Description and candidate experience.
- Do not invent skills unsupported by the Job Description or Experience Repository.

Mandatory ATS gap analysis requirements:
- Always identify Matched Keywords.
- Always identify Partial Match Keywords with 60-80% estimated coverage.
- Always identify Missing Important Keywords.
- Always consider partial matches and missing keywords during CV creation.
- Always list Partial Match Keywords and Missing Important Keywords in the Final ATS Report.
- The objective is zero deviations.
"""


ADDITIONAL_ATS_SKILLS_RULES = """
ADDITIONAL ATS SKILLS RULES - NON NEGOTIABLE

Additional ATS Skills must always be displayed vertically:
- Exactly 10 skills.
- One skill per line.
- No bullets.
- No numbering.
- No commas separating multiple skills.
- No blank lines between skills.
- No extra explanation before or after the skill list.
- Skills must be extracted from the Job Description and supported by candidate experience.
"""

ATS_GAP_ANALYSIS_RULES = """
MANDATORY ATS GAP ANALYSIS - NON NEGOTIABLE

The final output must always include ATS gap analysis at the end, after the Cover Letter when Cover Letter is requested.

The final ATS section must include these Markdown tables:

1. Keyword Coverage Table
| Category | Keyword | Status |
|---|---|---|
| Matched | keyword | Covered |
| Partial Match | keyword | 60-80% |
| Missing Important | keyword | High Priority |

2. Partial Match Keywords Table
| Partial Match Keyword | Estimated Coverage | Action Taken |
|---|---|---|
| keyword | 60-80% | Improved naturally in CV when supported by source facts |

3. Missing Important Keywords Table
| Missing Important Keyword | Priority | Action |
|---|---|---|
| keyword | High | Address if supported by candidate experience |

Rules:
- Always include Partial Match Keywords.
- Always include Missing Important Keywords.
- Partial matches are keywords with estimated 60-80% coverage.
- Missing Important Keywords must be extracted from the Job Description.
- Never invent keywords.
- Every partial match and missing important keyword must be considered during CV creation.
- The objective is zero ATS deviations.
- If a keyword cannot be added because it is not supported by the Experience Repository, keep it listed as Missing Important.
- Use Markdown tables so DOCX and PDF exports render real table cells.
"""

LANGUAGE_ALIGNMENT_RULES = """
LANGUAGE ALIGNMENT RULES

- Detect the language of the Job Description.
- The Cover Letter must always be written in the same language as the Job Description.
- If the Job Description is English, the Cover Letter must be English.
- If the Job Description is Dutch, the Cover Letter must be Dutch.
- If the Job Description is German, the Cover Letter must be German.
- If the Job Description is French, the Cover Letter must be French.
- If the Job Description is Portuguese, the Cover Letter must be Portuguese.
- Never translate the Cover Letter into another language unless the user explicitly selected a different output language.
"""

FINAL_OUTPUT_POLICY = """
FINAL OUTPUT POLICY

Do all analysis, scoring, ranking and validation internally.
Do not show reasoning, scoring steps, internal ranking, draft alternatives, keyword dumps or validation logs.

Final answer must contain only:
1. SELECTED CV TYPE
2. ATS VALIDATION STATUS - concise Markdown table with columns Metric and Score
3. FINAL TAILORED CV
4. FINAL ATS REPORT - concise realistic simulated scores, using Markdown tables for metrics

If Cover Letter is requested, include it inside FINAL TAILORED CV after ATS Analysis.

At the end of FINAL ATS REPORT, always include:
- Keyword Coverage Table
- Partial Match Keywords Table
- Missing Important Keywords Table

All tables must use Markdown table syntax so DOCX and PDF exports render formatted cells.
"""


def load_history() -> List[Dict]:
    if not HISTORY_FILE.exists():
        return []

    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as file:
            data = json.load(file)

        if isinstance(data, list):
            return data

        return []

    except Exception:
        return []


def save_history(history: List[Dict]) -> None:
    try:
        history = history[:MAX_HISTORY_ITEMS]

        with open(HISTORY_FILE, "w", encoding="utf-8") as file:
            json.dump(history, file, ensure_ascii=False, indent=2)

    except Exception:
        pass


def add_history_item(kind: str, title: str, job_description: str, result: str, model: str, context_profile: str) -> None:
    history = load_history()

    item = {
        "id": datetime.now().strftime("%Y%m%d_%H%M%S"),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "kind": kind,
        "title": title[:120] if title else kind,
        "job_description": job_description or "",
        "result": result or "",
        "model": model or "",
        "context_profile": context_profile or "",
    }

    history.insert(0, item)
    save_history(history)


def get_history_title(job_description: str, fallback: str = "CareerOps Output") -> str:
    if not job_description:
        return fallback

    clean_lines = [line.strip() for line in job_description.splitlines() if line.strip()]

    for line in clean_lines[:8]:
        if len(line) >= 8:
            return line[:100]

    return fallback


def render_history_panel() -> None:
    st.markdown("### 📚 Histórico")

    history = load_history()

    if not history:
        st.caption("Nenhuma saída salva ainda.")
        return

    labels = [
        f"{item.get('created_at', '')} · {item.get('kind', 'Output')} · {item.get('title', 'Untitled')}"
        for item in history
    ]

    selected_index = st.selectbox(
        "Saídas salvas",
        range(len(labels)),
        format_func=lambda index: labels[index],
        key="history_selected_index",
    )

    selected_item = history[selected_index]

    c1, c2 = st.columns(2)

    with c1:
        if st.button("📂 Carregar", use_container_width=True, key="load_btn"):
            st.session_state["last_result"] = selected_item.get("result", "")
            st.session_state["last_job"] = selected_item.get("job_description", "")
            st.session_state["loaded_history_item"] = selected_item
            st.success("Item do histórico carregado.")

    with c2:
        if st.button("🗑️ Limpar", use_container_width=True, key="clear_btn"):
            save_history([])
            st.session_state.pop("loaded_history_item", None)
            st.success("Histórico limpo.")
            st.rerun()

    with st.expander("👁️ Visualizar selecionado"):
        st.caption(f"Modelo: {selected_item.get('model', '')}")
        st.caption(f"Contexto: {selected_item.get('context_profile', '')}")
        st.text((selected_item.get("result", "") or "")[:2500])

def inject_css():
    st.markdown(
        """
        <style>
        :root {
            --bg-primary: #0f0f15;
            --bg-secondary: #1a1a24;
            --bg-tertiary: #25252f;
            --border-color: rgba(255, 255, 255, 0.08);
            --border-color-hover: rgba(255, 255, 255, 0.12);
            --text-primary: #ffffff;
            --text-secondary: #b4b4b8;
            --text-tertiary: #8b8b92;
            --accent-primary: #10a37f;
            --accent-secondary: #8b5cf6;
            --accent-tertiary: #ec4899;
            --success: #10a37f;
            --warning: #f59e0b;
            --error: #ef4444;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        html, body, [data-testid="stAppViewContainer"] {
            background-color: var(--bg-primary) !important;
            color: var(--text-primary) !important;
        }

        .stApp {
            background: linear-gradient(135deg, var(--bg-primary) 0%, var(--bg-secondary) 50%, var(--bg-primary) 100%);
            color: var(--text-primary);
        }

        [data-testid="stHeader"] {
            background: transparent !important;
        }

        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, var(--bg-secondary) 0%, var(--bg-tertiary) 100%) !important;
            border-right: 1px solid var(--border-color) !important;
        }

        [data-testid="stSidebar"] * {
            color: var(--text-primary) !important;
        }

        .block-container {
            max-width: 1200px;
            padding: 2rem 1.5rem;
        }

        /* Header & Title Styles */
        .header-container {
            background: linear-gradient(135deg, rgba(16, 163, 127, 0.1) 0%, rgba(139, 92, 246, 0.05) 100%);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 2rem;
            margin-bottom: 2rem;
            backdrop-filter: blur(10px);
        }

        .header-top {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            gap: 1.5rem;
            margin-bottom: 1.5rem;
        }

        .header-brand {
            font-size: 0.75rem;
            font-weight: 700;
            letter-spacing: 0.15em;
            text-transform: uppercase;
            color: var(--accent-primary);
            margin-bottom: 0.5rem;
        }

        .header-title {
            font-size: 2.5rem;
            font-weight: 800;
            line-height: 1.1;
            letter-spacing: -0.02em;
            color: var(--text-primary);
            margin-bottom: 0.5rem;
        }

        .header-subtitle {
            font-size: 1rem;
            color: var(--text-secondary);
            line-height: 1.5;
            max-width: 700px;
        }

        .status-badge {
            background: rgba(16, 163, 127, 0.15);
            border: 1px solid var(--accent-primary);
            border-radius: 999px;
            padding: 0.5rem 1rem;
            font-size: 0.85rem;
            color: var(--accent-primary);
            white-space: nowrap;
            font-weight: 600;
        }

        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin-top: 1.5rem;
        }

        .stat-card {
            background: rgba(255, 255, 255, 0.03);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 1rem;
            transition: all 0.3s ease;
        }

        .stat-card:hover {
            background: rgba(255, 255, 255, 0.05);
            border-color: var(--border-color-hover);
        }

        .stat-label {
            font-size: 0.75rem;
            text-transform: uppercase;
            letter-spacing: 0.1em;
            color: var(--text-tertiary);
            margin-bottom: 0.5rem;
            font-weight: 600;
        }

        .stat-value {
            font-size: 1.25rem;
            font-weight: 700;
            color: var(--accent-primary);
        }

        /* Card Styles */
        .card {
            background: linear-gradient(135deg, rgba(255, 255, 255, 0.03) 0%, rgba(255, 255, 255, 0.01) 100%);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 2rem;
            margin-bottom: 1.5rem;
            backdrop-filter: blur(10px);
            transition: all 0.3s ease;
        }

        .card:hover {
            border-color: var(--border-color-hover);
            background: linear-gradient(135deg, rgba(255, 255, 255, 0.05) 0%, rgba(255, 255, 255, 0.02) 100%);
        }

        .card-title {
            font-size: 1.5rem;
            font-weight: 700;
            color: var(--text-primary);
            margin-bottom: 0.5rem;
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }

        .card-subtitle {
            font-size: 0.95rem;
            color: var(--text-secondary);
            line-height: 1.5;
            margin-bottom: 1.5rem;
        }

        /* Input Styles */
        .stTextArea textarea,
        .stTextInput input,
        .stSelectbox select {
            background-color: rgba(255, 255, 255, 0.05) !important;
            border: 1px solid var(--border-color) !important;
            border-radius: 12px !important;
            color: var(--text-primary) !important;
            padding: 0.75rem !important;
            font-size: 0.95rem !important;
            transition: all 0.3s ease !important;
        }

        .stTextArea textarea:focus,
        .stTextInput input:focus,
        .stSelectbox select:focus {
            background-color: rgba(255, 255, 255, 0.08) !important;
            border-color: var(--accent-primary) !important;
            box-shadow: 0 0 0 3px rgba(16, 163, 127, 0.1) !important;
        }

        /* Button Styles */
        .stButton > button,
        .stDownloadButton > button {
            background: linear-gradient(135deg, var(--accent-primary) 0%, var(--accent-secondary) 100%) !important;
            color: white !important;
            border: 0 !important;
            border-radius: 12px !important;
            padding: 0.75rem 1.5rem !important;
            font-weight: 600 !important;
            font-size: 0.95rem !important;
            transition: all 0.3s ease !important;
            cursor: pointer !important;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 8px 24px rgba(16, 163, 127, 0.3) !important;
        }

        .stButton > button:active,
        .stDownloadButton > button:active {
            transform: translateY(0) !important;
        }

        /* Secondary Button */
        .stButton > button[kind="secondary"] {
            background: rgba(255, 255, 255, 0.1) !important;
            border: 1px solid var(--border-color) !important;
            color: var(--text-primary) !important;
        }

        .stButton > button[kind="secondary"]:hover {
            background: rgba(255, 255, 255, 0.15) !important;
            border-color: var(--border-color-hover) !important;
        }

        /* Tabs */
        [data-testid="stTabs"] button {
            border-radius: 12px !important;
            padding: 0.75rem 1.25rem !important;
            background: transparent !important;
            color: var(--text-secondary) !important;
            border: 0 !important;
            font-weight: 600 !important;
            transition: all 0.3s ease !important;
        }

        [data-testid="stTabs"] button[aria-selected="true"] {
            color: var(--accent-primary) !important;
            background: rgba(16, 163, 127, 0.1) !important;
            border-bottom: 2px solid var(--accent-primary) !important;
        }

        /* Expander */
        .streamlit-expanderHeader {
            background: rgba(255, 255, 255, 0.03) !important;
            border: 1px solid var(--border-color) !important;
            border-radius: 12px !important;
            padding: 1rem !important;
        }

        .streamlit-expanderHeader:hover {
            background: rgba(255, 255, 255, 0.05) !important;
        }

        /* Messages */
        .stSuccess, .stInfo, .stWarning, .stError {
            border-radius: 12px !important;
            padding: 1rem !important;
            border: 1px solid !important;
        }

        .stSuccess {
            background: rgba(16, 163, 127, 0.1) !important;
            border-color: var(--success) !important;
            color: var(--success) !important;
        }

        .stInfo {
            background: rgba(59, 130, 246, 0.1) !important;
            border-color: #3b82f6 !important;
            color: #3b82f6 !important;
        }

        .stWarning {
            background: rgba(245, 158, 11, 0.1) !important;
            border-color: var(--warning) !important;
            color: var(--warning) !important;
        }

        .stError {
            background: rgba(239, 68, 68, 0.1) !important;
            border-color: var(--error) !important;
            color: var(--error) !important;
        }

        /* Result Shell */
        .result-container {
            background: rgba(255, 255, 255, 0.02);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 2rem;
            margin-top: 1.5rem;
            line-height: 1.7;
            color: var(--text-primary);
        }

        .result-container h1,
        .result-container h2,
        .result-container h3 {
            color: var(--accent-primary);
            margin-top: 1.5rem;
            margin-bottom: 0.75rem;
        }

        .result-container h1 {
            font-size: 1.75rem;
        }

        .result-container h2 {
            font-size: 1.35rem;
        }

        .result-container h3 {
            font-size: 1.1rem;
        }

        .result-container table {
            width: 100%;
            border-collapse: collapse;
            margin: 1rem 0;
            border-radius: 8px;
            overflow: hidden;
        }

        .result-container table th {
            background: rgba(16, 163, 127, 0.15);
            color: var(--accent-primary);
            padding: 0.75rem;
            text-align: left;
            font-weight: 600;
            border-bottom: 2px solid var(--border-color);
        }

        .result-container table td {
            padding: 0.75rem;
            border-bottom: 1px solid var(--border-color);
            color: var(--text-secondary);
        }

        .result-container table tr:hover {
            background: rgba(255, 255, 255, 0.02);
        }

        /* Audit Box */
        .audit-box {
            background: rgba(16, 163, 127, 0.08);
            border: 1px solid rgba(16, 163, 127, 0.3);
            border-radius: 12px;
            padding: 1rem;
            margin: 1rem 0;
            color: var(--text-secondary);
            font-size: 0.9rem;
            line-height: 1.6;
        }

        /* File Uploader */
        [data-testid="stFileUploader"] {
            border: 2px dashed var(--border-color) !important;
            background: rgba(255, 255, 255, 0.02) !important;
            border-radius: 12px !important;
            padding: 1.5rem !important;
            transition: all 0.3s ease !important;
        }

        [data-testid="stFileUploader"]:hover {
            border-color: var(--accent-primary) !important;
            background: rgba(16, 163, 127, 0.05) !important;
        }

        /* Divider */
        hr {
            border-color: var(--border-color) !important;
            margin: 1.5rem 0 !important;
        }

        /* Sidebar Styles */
        .sidebar-section {
            margin-bottom: 2rem;
        }

        .sidebar-section-title {
            font-size: 0.85rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.1em;
            color: var(--text-tertiary);
            margin-bottom: 1rem;
            margin-top: 1.5rem;
        }

        .sidebar-section:first-child .sidebar-section-title {
            margin-top: 0;
        }

        /* Footer */
        .footer {
            text-align: center;
            color: var(--text-tertiary);
            font-size: 0.85rem;
            padding: 2rem 0;
            margin-top: 3rem;
            border-top: 1px solid var(--border-color);
        }

        /* Responsive */
        @media (max-width: 768px) {
            .header-title {
                font-size: 1.75rem;
            }

            .header-top {
                flex-direction: column;
            }

            .stats-grid {
                grid-template-columns: repeat(2, 1fr);
            }

            .card {
                padding: 1.5rem;
            }
        }

        /* Scrollbar Styling */
        ::-webkit-scrollbar {
            width: 8px;
            height: 8px;
        }

        ::-webkit-scrollbar-track {
            background: transparent;
        }

        ::-webkit-scrollbar-thumb {
            background: rgba(255, 255, 255, 0.1);
            border-radius: 4px;
        }

        ::-webkit-scrollbar-thumb:hover {
            background: rgba(255, 255, 255, 0.2);
        }

        /* Metric Styles */
        [data-testid="metric-container"] {
            background: rgba(255, 255, 255, 0.03) !important;
            border: 1px solid var(--border-color) !important;
            border-radius: 12px !important;
            padding: 1.25rem !important;
        }

        /* Caption */
        .stCaption {
            color: var(--text-tertiary) !important;
        }

        </style>
        """,
        unsafe_allow_html=True,
    )


def load_template_text(filename: str, fallback: str = "") -> str:
    path = TEMPLATE_DIR / filename
    if not path.exists():
        return fallback
    try:
        if path.suffix.lower() == ".docx":
            doc = Document(path)
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return fallback


def load_first_available(filenames: List[str]) -> str:
    for filename in filenames:
        text = load_template_text(filename)
        if text.strip():
            return text
    return ""


def get_clients():
    keys = [
        os.getenv("GEMINI_API_KEY_1"),
        os.getenv("GEMINI_API_KEY_2"),
        os.getenv("GEMINI_API_KEY_3"),
        os.getenv("GEMINI_API_KEY_4"),
        os.getenv("GEMINI_API_KEY_5"),
        os.getenv("GEMINI_API_KEY"),
    ]

    clients = []

    for index, key in enumerate(keys, start=1):
        if key and key.strip():
            clients.append({
                "label": f"API Key {index}",
                "client": genai.Client(api_key=key.strip()),
            })

    if not clients:
        st.error(
            "No Gemini API key found. Add GEMINI_API_KEY_1=your_key_here to the .env file. "
            "You can also add GEMINI_API_KEY_2, GEMINI_API_KEY_3, GEMINI_API_KEY_4 and GEMINI_API_KEY_5."
        )
        st.stop()

    return clients


def read_uploaded_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    uploaded_file.seek(0)
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if name.endswith((".txt", ".md", ".csv")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        doc = Document(BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if name.endswith(".pdf"):
        if PdfReader is None:
            return "[PDF uploaded, but pypdf is not installed.]"
        reader = PdfReader(BytesIO(data))
        return "\n".join([(page.extract_text() or "") for page in reader.pages])
    return data.decode("utf-8", errors="ignore")


def truncate_text(text: str, max_chars: int) -> str:
    text = text or ""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[Content truncated by application limit]"


def is_markdown_table_line(line: str) -> bool:
    stripped = (line or "").strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_markdown_separator_line(line: str) -> bool:
    stripped = (line or "").strip()
    if not is_markdown_table_line(stripped):
        return False

    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_markdown_table(lines: List[str], start_index: int) -> Tuple[List[List[str]], int]:
    table_lines = []
    index = start_index

    while index < len(lines) and is_markdown_table_line(lines[index]):
        table_lines.append(lines[index])
        index += 1

    rows = []

    for raw in table_lines:
        if is_markdown_separator_line(raw):
            continue

        cells = [cell.strip() for cell in raw.strip().strip("|").split("|")]
        rows.append(cells)

    return rows, index


def detect_pipe_table_block(lines: List[str], start_index: int) -> bool:
    if start_index >= len(lines):
        return False

    if not is_markdown_table_line(lines[start_index]):
        return False

    if start_index + 1 < len(lines) and is_markdown_separator_line(lines[start_index + 1]):
        return True

    return False


def add_docx_markdown_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]

    table = doc.add_table(rows=len(normalized_rows), cols=max_cols)
    table.style = "Table Grid"

    for row_index, row in enumerate(normalized_rows):
        for col_index, cell_text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = cell_text

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    if row_index == 0:
                        run.bold = True

    doc.add_paragraph("")


def add_pdf_markdown_table(story: List, rows: List[List[str]], styles) -> None:
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]

    table_data = []

    for row in normalized_rows:
        pdf_row = []
        for cell in row:
            safe = str(cell).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            pdf_row.append(Paragraph(safe, styles["TableCell"]))
        table_data.append(pdf_row)

    page_width = A4[0] - (1.65 * cm * 2)
    col_width = page_width / max_cols
    table = Table(table_data, colWidths=[col_width] * max_cols, repeatRows=1)

    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.55, colors.HexColor("#8A94A6")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#101828")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#FFFFFF")),
        ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#182230")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#FFFFFF"), colors.HexColor("#F8FAFC")]),
    ]))

    story.append(table)
    story.append(Spacer(1, 0.22 * cm))

def make_docx(content: str, title: str = "CareerOps Studio Output") -> bytes:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"

    doc.add_heading(title, level=1)

    lines = content.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            index += 1
            continue

        if detect_pipe_table_block(lines, index):
            rows, next_index = parse_markdown_table(lines, index)
            add_docx_markdown_table(doc, rows)
            index = next_index
            continue

        upper = stripped.upper()

        if upper in [
            "SELECTED CV TYPE",
            "ATS VALIDATION STATUS",
            "FINAL TAILORED CV",
            "FINAL ATS REPORT",
            "COVER LETTER",
        ]:
            doc.add_heading(stripped, level=1)
        elif stripped in REQUIRED_CV_SECTIONS or stripped in ROLE_NAMES:
            doc.add_heading(stripped, level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        else:
            doc.add_paragraph(stripped)

        index += 1

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def make_pdf(content: str, title: str = "CareerOps Studio Output") -> bytes:
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        rightMargin=1.65 * cm,
        leftMargin=1.65 * cm,
        topMargin=1.55 * cm,
        bottomMargin=1.55 * cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ExecutiveTitle", parent=styles["Title"], alignment=TA_LEFT, fontSize=17, leading=21, spaceAfter=12))
    styles.add(ParagraphStyle(name="ExecutiveHeading1", parent=styles["Heading1"], alignment=TA_LEFT, fontSize=14, leading=17, spaceBefore=10, spaceAfter=7, textColor=colors.HexColor("#1F4E79")))
    styles.add(ParagraphStyle(name="ExecutiveHeading2", parent=styles["Heading2"], alignment=TA_LEFT, fontSize=11.5, leading=14, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor("#101828")))
    styles.add(ParagraphStyle(name="ExecutiveNormal", parent=styles["Normal"], fontSize=9.3, leading=12.2, spaceAfter=4.8))
    styles.add(ParagraphStyle(name="TableCell", parent=styles["Normal"], fontSize=8.5, leading=10.5, spaceAfter=0))

    story = [Paragraph(title, styles["ExecutiveTitle"])]

    lines = content.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if not stripped:
            story.append(Spacer(1, 0.18 * cm))
            index += 1
            continue

        if detect_pipe_table_block(lines, index):
            rows, next_index = parse_markdown_table(lines, index)
            add_pdf_markdown_table(story, rows, styles)
            index = next_index
            continue

        safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        upper = stripped.upper()

        if upper in [
            "SELECTED CV TYPE",
            "ATS VALIDATION STATUS",
            "FINAL TAILORED CV",
            "FINAL ATS REPORT",
            "COVER LETTER",
        ]:
            story.append(Paragraph(safe, styles["ExecutiveHeading1"]))
        elif stripped in REQUIRED_CV_SECTIONS or stripped in ROLE_NAMES:
            story.append(Paragraph(safe, styles["ExecutiveHeading2"]))
        else:
            story.append(Paragraph(safe, styles["ExecutiveNormal"]))

        index += 1

    pdf.build(story)
    return bio.getvalue()

def clean_model_text(text: str) -> str:
    if not text:
        return ""
    markers = ["done thinking.", "</think>"]
    lower = text.lower()
    for marker in markers:
        idx = lower.rfind(marker)
        if idx != -1:
            text = text[idx + len(marker):]
            break
    return text.strip()


def call_model(prompt: str, primary_model: str, failover_models: List[str]) -> Tuple[str, str, List[str]]:
    clients = get_clients()

    models_to_try = []
    for item in [primary_model] + failover_models:
        item = (item or "").strip()
        if item and item not in models_to_try:
            models_to_try.append(item)

    if len(prompt) > 180000:
        prompt = prompt[:180000] + "\n\n[Prompt truncated by application limit]"

    attempts = []
    last_error = None

    for client_info in clients:
        client_label = client_info["label"]
        client = client_info["client"]

        for current_model in models_to_try:
            for attempt in range(2):
                try:
                    attempts.append(f"{client_label} / {current_model}: attempt {attempt + 1}")

                    response = client.models.generate_content(
                        model=current_model,
                        contents=prompt,
                    )

                    text = clean_model_text(response.text or "")

                    if not text:
                        raise RuntimeError("Empty response returned by provider.")

                    return text, f"{current_model} ({client_label})", attempts

                except Exception as exc:
                    last_error = exc
                    error_text = str(exc).lower()
                    attempts.append(f"FAILED: {client_label} / {current_model}")

                    if any(code in error_text for code in ["503", "unavailable", "high demand", "overloaded", "500"]):
                        time.sleep(2 + attempt * 3)
                        continue

                    if any(code in error_text for code in ["429", "quota", "rate", "resource_exhausted"]):
                        break

                    if any(code in error_text for code in ["404", "not found", "invalid", "model"]):
                        break

                    break

    raise RuntimeError(
        "All configured API keys and models failed. Last error: "
        + str(last_error)
        + "\n\nLast attempts:\n"
        + "\n".join(attempts[-20:])
    )

def normalize_role_name(line: str) -> str:
    cleaned = re.sub(r"^[#\-\*\•\s\d\.\)]+", "", line or "").strip()
    cleaned = cleaned.replace(":", "").strip()
    return cleaned


def normalize_for_match(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()


def detect_cv_type(result: str) -> str:
    text = result or ""
    lower = text.lower()

    m = re.search(r"selected\s+cv\s+type\s*[:\n\-]*\s*(manager cv|director cv)", lower)
    if m:
        return "Manager CV" if "manager" in m.group(1) else "Director CV"

    first_500 = lower[:500]
    if "director cv" in first_500:
        return "Director CV"
    if "manager cv" in first_500:
        return "Manager CV"

    return "Unknown"


def is_section_heading(line: str) -> bool:
    normalized = normalize_for_match(line)
    known = [
        "selected cv type",
        "ats validation status",
        "final tailored cv",
        "name",
        "contact details",
        "additional ats skills",
        "professional summary",
        "professional experience",
        "education",
        "systems tools",
        "ats analysis",
        "cover letter",
        "final ats report",
    ]
    return normalized in known


def is_role_heading(line: str) -> str:
    normalized = normalize_role_name(line).lower()
    for role in ROLE_NAMES:
        if normalized == role.lower():
            return role
    return ""


def is_plain_achievement_line(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped:
        return False
    if is_section_heading(stripped):
        return False
    if is_role_heading(stripped):
        return False
    forbidden_prefixes = ("achievement:", "achievements:", "bullet:", "result:", "success:", "kpi:")
    if stripped.lower().startswith(forbidden_prefixes):
        return False
    if stripped.startswith(("-", "*", "•")):
        return False
    if re.match(r"^\d+[\.\)]\s+", stripped):
        return False
    if len(stripped) < 80:
        return False
    return True


def extract_role_achievements(result: str) -> Dict[str, List[str]]:
    role_hits = {role: [] for role in ROLE_NAMES}
    current_role = None
    inside_experience = False

    for raw_line in (result or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        normalized_section = normalize_for_match(line)

        if normalized_section == "professional experience":
            inside_experience = True
            current_role = None
            continue

        if inside_experience and normalized_section in ["education", "systems tools", "ats analysis", "cover letter", "final ats report"]:
            break

        role = is_role_heading(line)
        if role:
            current_role = role
            inside_experience = True
            continue

        if inside_experience and current_role and is_plain_achievement_line(line):
            role_hits[current_role].append(line)

    return role_hits


def extract_section_text(result: str, section_name: str, stop_sections: List[str]) -> str:
    text = result or ""
    lines = text.splitlines()
    collecting = False
    collected = []
    wanted = normalize_for_match(section_name)
    stops = {normalize_for_match(s) for s in stop_sections}

    for line in lines:
        normalized = normalize_for_match(line.strip())
        if normalized == wanted:
            collecting = True
            continue
        if collecting and normalized in stops:
            break
        if collecting:
            collected.append(line)

    return "\n".join(collected).strip()


def extract_cover_letter_text(result: str) -> str:
    return extract_section_text(result, "Cover Letter", ["Final ATS Report"])


def extract_additional_ats_skills(result: str) -> List[str]:
    section = extract_section_text(
        result,
        "Additional ATS Skills",
        ["Professional Summary", "Professional Experience", "Education", "Systems / Tools", "ATS Analysis", "Cover Letter", "Final ATS Report"],
    )
    skills = []
    for line in section.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        stripped = re.sub(r"^[\-\*\•\d\.\)\s]+", "", stripped).strip()
        if stripped:
            skills.append(stripped)
    return skills




def extract_additional_ats_skills_raw_section(result: str) -> str:
    return extract_section_text(
        result,
        "Additional ATS Skills",
        ["Professional Summary", "Professional Experience", "Education", "Systems / Tools", "ATS Analysis", "Cover Letter", "Final ATS Report"],
    )


def validate_vertical_ats_skills(result: str) -> List[str]:
    issues = []
    section = extract_additional_ats_skills_raw_section(result)
    skills = extract_additional_ats_skills(result)

    if not section.strip():
        issues.append("Additional ATS Skills section is missing or empty.")
        return issues

    raw_lines = section.splitlines()
    if any(not line.strip() for line in raw_lines):
        issues.append("Additional ATS Skills contains blank lines. Use one skill per line with no blank lines.")

    for index, line in enumerate(raw_lines, start=1):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("-", "*", "•")) or re.match(r"^\d+[\.\)]\s+", stripped):
            issues.append(f"Additional ATS Skills line {index}: bullets or numbering are not allowed.")
        if "," in stripped:
            issues.append(f"Additional ATS Skills line {index}: comma-separated skills are not allowed. Use one skill per line.")

    if len(skills) != 10:
        issues.append(f"Additional ATS Skills: expected exactly 10 skills; found {len(skills)}.")

    return issues


def count_blank_lines_between_role_achievements(result: str) -> List[str]:
    issues = []
    lines = (result or "").splitlines()
    inside_experience = False
    current_role = None
    previous_was_achievement = False

    for raw_line in lines:
        line = raw_line.strip()
        normalized = normalize_for_match(line)

        if normalized == "professional experience":
            inside_experience = True
            current_role = None
            previous_was_achievement = False
            continue

        if inside_experience and normalized in ["education", "systems tools", "ats analysis", "cover letter", "final ats report"]:
            break

        role = is_role_heading(line)
        if role:
            current_role = role
            previous_was_achievement = False
            continue

        if inside_experience and current_role:
            if not line and previous_was_achievement:
                issues.append(f"{current_role}: blank line found between achievements. Achievements must be compact with no blank lines.")
                previous_was_achievement = False
                continue

            if is_plain_achievement_line(line):
                previous_was_achievement = True
            elif line:
                previous_was_achievement = False

    return issues


def has_markdown_table_with_header(result: str, required_headers: List[str]) -> bool:
    lines = (result or "").splitlines()
    required = [h.lower().strip() for h in required_headers]

    for line in lines:
        if not is_markdown_table_line(line):
            continue
        cells = [cell.strip().lower() for cell in line.strip().strip("|").split("|")]
        if all(header in cells for header in required):
            return True

    return False


def validate_ats_gap_sections(result: str) -> List[str]:
    issues = []
    lower = (result or "").lower()

    for term in ["keyword coverage", "partial match", "missing important"]:
        if term not in lower:
            issues.append(f"Final ATS Report missing required ATS gap section or label: {term}.")

    if not has_markdown_table_with_header(result, ["Category", "Keyword"]):
        issues.append("Keyword Coverage Table missing. Required Markdown table with columns Category and Keyword.")

    if not has_markdown_table_with_header(result, ["Partial Match Keyword", "Estimated Coverage"]):
        issues.append("Partial Match Keywords Table missing. Required Markdown table with columns Partial Match Keyword and Estimated Coverage.")

    if not has_markdown_table_with_header(result, ["Missing Important Keyword", "Priority"]):
        issues.append("Missing Important Keywords Table missing. Required Markdown table with columns Missing Important Keyword and Priority.")

    return issues


def validate_cover_letter_language_instruction(result: str, output_scope: str) -> List[str]:
    issues = []
    if not has_requested_cover_letter(output_scope):
        return issues
    cover = extract_cover_letter_text(result)
    if cover and ("same language" in cover.lower() or "job description language" in cover.lower()):
        issues.append("Cover Letter contains language instruction text instead of a real letter.")
    return issues


def has_requested_cover_letter(output_scope: str) -> bool:
    return "cover letter" in (output_scope or "").lower()


def validate_required_sections(result: str, output_scope: str) -> List[str]:
    issues = []
    lower_normalized_lines = {normalize_for_match(line.strip()) for line in (result or "").splitlines() if line.strip()}

    for section in REQUIRED_CV_SECTIONS:
        if normalize_for_match(section) not in lower_normalized_lines:
            issues.append(f"Missing required CV section: {section}.")

    if has_requested_cover_letter(output_scope) and normalize_for_match("Cover Letter") not in lower_normalized_lines:
        issues.append("Cover Letter requested but Cover Letter section is missing.")

    return issues


def validate_generated_output(result: str, output_scope: str) -> Tuple[bool, List[str]]:
    issues = []
    issues.extend(validate_required_sections(result, output_scope))

    cv_type = detect_cv_type(result)
    if cv_type not in EXPECTED_COUNTS:
        issues.append("SELECTED CV TYPE missing or unclear. Must be exactly Manager CV or Director CV.")
        return False, issues

    role_achievements = extract_role_achievements(result)
    expected = EXPECTED_COUNTS[cv_type]

    for role, required_count in expected.items():
        actual_count = len(role_achievements.get(role, []))
        if required_count == 0 and actual_count > 0:
            issues.append(f"{role}: must not appear in {cv_type}; found {actual_count} achievement lines.")
        elif required_count > 0 and actual_count != required_count:
            issues.append(f"{role}: expected exactly {required_count} achievement lines; found {actual_count}.")

        for idx, achievement in enumerate(role_achievements.get(role, []), start=1):
            char_count = len(achievement)
            if achievement.lower().startswith(("achievement:", "achievements:", "bullet:", "result:", "success:", "kpi:")):
                issues.append(f"{role} achievement {idx}: forbidden label/prefix found. Achievements must be plain text only.")
            if achievement.startswith(("-", "*", "•")) or re.match(r"^\d+[\.\)]\s+", achievement):
                issues.append(f"{role} achievement {idx}: bullet symbol or numbered list found. Use plain text only.")
            if char_count < 170 or char_count > 190:
                issues.append(f"{role} achievement {idx}: {char_count} characters. Required 170-190 including spaces.")

    issues.extend(count_blank_lines_between_role_achievements(result))
    issues.extend(validate_vertical_ats_skills(result))
    issues.extend(validate_ats_gap_sections(result))
    issues.extend(validate_cover_letter_language_instruction(result, output_scope))

    if has_requested_cover_letter(output_scope):
        cover = extract_cover_letter_text(result)
        if not cover:
            issues.append("Cover Letter requested but Cover Letter section was not found.")
        else:
            cover_len = len(cover)
            if cover_len < 1795 or cover_len > 1805:
                issues.append(f"Cover Letter length is {cover_len} characters. Required 1,795-1,805; target 1,800.")

    lower = (result or "").lower()
    external_terms = ["linkedin ats", "jobscan", "skillsyncer", "resume worded", "rezi"]
    if any(term in lower for term in external_terms):
        if not any(term in lower for term in ["simulated", "simulation", "internal", "ats-style", "style"]):
            issues.append("ATS validation references external tools but does not state that scores are internal simulations.")

    return len(issues) == 0, issues

def build_repair_prompt(original_prompt: str, previous_result: str, issues: List[str], output_scope: str) -> str:
    issue_text = "\n".join(f"- {issue}" for issue in issues)

    return f"""
Repair the previous output. Return only the corrected final output. Do not explain.

STRICT VALIDATION FAILURES:
{issue_text}

NON-NEGOTIABLE REPAIR RULES:
- Keep facts faithful to the Experience Repository.
- Keep the selected CV type unless it is missing or clearly wrong.
- Use the mandatory CV structure and section order.
- Use exact role headings.
- Each achievement must be plain text on its own line.
- Do not write ACHIEVEMENT:, Achievement:, Bullet:, Result:, Success:, KPI: or any similar label.
- Do not use bullet symbols, numbered lists or multiple achievements in one paragraph.
- Every achievement line must be 170-190 characters including spaces.
- No blank lines between achievements inside the same role.
- Manager CV must have exactly 8 Trade Compliance Manager EMEA achievements. Not 7.
- Manager CV must have exactly 7 Head of Logistics EMEA achievements. Not 6.
- Director CV must not include Trade Compliance Manager EMEA.
- Additional ATS Skills must contain exactly 10 skills, one per line, no bullets, no numbering and no blank lines.
- Cover Letter, if requested, must be 1,795-1,805 characters including spaces, target 1,800.
- Cover Letter must match the language of the Job Description unless the user explicitly requested another language.
- ATS validation must be described as simulated/internal, not live website validation.
- Final ATS Report must include Keyword Coverage Table, Partial Match Keywords Table and Missing Important Keywords Table.
- Partial matches must be estimated 60-80% coverage.
- Missing Important Keywords must be extracted from the Job Description.
- The objective is zero ATS deviations.
- Final content must be easy to copy and paste into the CV template with compact ATS-friendly formatting.

ORIGINAL TASK:
{original_prompt}

PREVIOUS OUTPUT:
{previous_result}

OUTPUT PACKAGE:
{output_scope}
"""


def generate_with_strict_validation(prompt: str, output_scope: str, primary_model: str, failover_models: List[str]) -> Tuple[str, str, List[str], List[str]]:
    result, used_model, attempts = call_model(prompt, primary_model, failover_models)
    is_valid, issues = validate_generated_output(result, output_scope)
    repair_round = 0

    while not is_valid and repair_round < 5:
        repair_round += 1
        repair_prompt = build_repair_prompt(prompt, result, issues, output_scope)
        repaired, repair_model, repair_attempts = call_model(repair_prompt, primary_model, failover_models)
        attempts.extend(repair_attempts)
        result = repaired
        used_model = repair_model
        is_valid, issues = validate_generated_output(result, output_scope)

    return result, used_model, attempts, issues

def build_context(instructions, master, manager, director, profile: str):
    limits = CONTEXT_PROFILES.get(profile, CONTEXT_PROFILES["Balanced"])
    return f"""
[OPERATING INSTRUCTIONS]
{truncate_text(instructions, limits['instructions'])}

[EXPERIENCE REPOSITORY - SOURCE OF TRUTH]
{truncate_text(master, limits['master'])}

[MANAGER CV TEMPLATE]
{truncate_text(manager, limits['manager'])}

[DIRECTOR CV TEMPLATE]
{truncate_text(director, limits['director'])}
"""


def render_downloads(result: str, base_name: str):
    st.markdown('<div class="card"><div class="card-title">📥 Exportar Resultado</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button("📄 TXT", result.encode("utf-8"), f"{base_name}.txt", "text/plain", use_container_width=True)
    with c2:
        st.download_button("📋 DOCX", make_docx(result), f"{base_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with c3:
        st.download_button("📑 PDF", make_pdf(result), f"{base_name}.pdf", "application/pdf", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)


def strict_cv_prompt(context: str, job_description: str, language: str, output_scope: str) -> str:
    return f"""
You are CareerOps Studio, an executive CV optimization engine. Do not describe yourself as AI in the output.

PRIMARY OBJECTIVE
Create a tailored CV from the Job Description using the Experience Repository as the only source of truth. The final output must be compact, ATS optimized and easy to copy and paste into the existing CV template.

ABSOLUTE RULES
1. Never invent achievements, metrics, companies, roles, dates, systems, certifications, countries, savings, headcount, budget or KPIs.
2. Select achievements only from the Experience Repository based on the Job Description.
3. Choose Manager CV or Director CV using the Job Description and instructions.
4. Follow exact role inclusion/exclusion and exact achievement counts.
5. Each achievement must be plain text on its own separate line.
6. Do not write ACHIEVEMENT:, Achievement:, Bullet:, Result:, Success:, KPI: or any similar label.
7. Do not use bullet symbols, numbered lists or multiple achievements in one paragraph.
8. Every achievement line must be 170-190 characters including spaces.
9. No blank lines between achievements inside the same role.
10. Manager CV must include exactly 8 Trade Compliance Manager EMEA achievement lines, not 7.
11. Manager CV must include exactly 7 Head of Logistics EMEA achievement lines, not 6.
12. Additional ATS Skills must contain exactly 10 skills extracted from the Job Description and supported by candidate experience.
13. Additional ATS Skills must be listed vertically: one skill per line, no bullets, no numbering, no blank lines.
14. Cover Letter, when requested, must be 1,795-1,805 characters including spaces. Target exactly 1,800.
15. Cover Letter must be written in the same language as the Job Description unless the user explicitly selected another output language.
16. ATS validation is an internal simulation. Never claim it is coming from live external websites.
17. Do all analysis internally. Do not show reasoning, scoring steps, keyword dumps, validation logs, draft alternatives or chain-of-thought.
18. Partial Match Keywords and Missing Important Keywords must be considered during CV creation and listed at the end.
19. The objective is zero ATS deviations.

{MANDATORY_CV_TYPE_RULES}

{MANDATORY_BULLET_COUNTS}

{STRICT_FORMAT_RULES}

{ADDITIONAL_ATS_SKILLS_RULES}

{COVER_LETTER_RULES}

{LANGUAGE_ALIGNMENT_RULES}

{ATS_RULES}

{ATS_GAP_ANALYSIS_RULES}

{FINAL_OUTPUT_POLICY}

MANDATORY CV STRUCTURE AND ORDER
Inside FINAL TAILORED CV, use this exact order:
Name
Contact Details
Additional ATS Skills
Professional Summary
Professional Experience
Education
Systems / Tools
ATS Analysis
Cover Letter only if requested

MANDATORY INTERNAL PROCESS
A. Extract company name, exact job title, seniority, location, industry, hard skills, soft skills, systems, certifications, leadership requirements and critical ATS keywords from the Job Description.
B. Select Manager CV or Director CV.
C. Apply mandatory role counts exactly.
D. Select and rewrite achievements from source facts only.
E. Validate each achievement line internally until it is 170-190 characters including spaces.
F. Place each achievement on a new line with no prefix and no bullet symbol.
G. Write Additional ATS Skills as exactly 10 vertical lines with no blank lines.
H. If Cover Letter is requested, write it from Job Description + selected achievements, match the Job Description language, and keep it 1,795-1,805 characters.
I. Produce internal simulated ATS-style statistics and skill matching as Markdown tables so PDF and DOCX exports render real table cells.
J. Identify matched keywords, partial matches at 60-80%, and missing important keywords from the Job Description.
K. Improve the CV naturally using supported partial/missing keywords; unsupported missing keywords must remain listed.
L. Return only the final output sections.

FINAL ATS REPORT REQUIRED TABLES
Always include these tables:
| Category | Keyword | Status |
|---|---|---|
| Matched | keyword | Covered |
| Partial Match | keyword | 60-80% |
| Missing Important | keyword | High Priority |

| Partial Match Keyword | Estimated Coverage | Action Taken |
|---|---|---|
| keyword | 60-80% | Improved naturally when supported |

| Missing Important Keyword | Priority | Action |
|---|---|---|
| keyword | High | Address if supported by source facts |

OUTPUT REQUIRED
{output_scope}

OUTPUT LANGUAGE
{language}

CONTEXT FILES
{context}

JOB DESCRIPTION
{job_description}
"""

def ats_validation_prompt(context: str, cv_text: str, job_description: str, language: str) -> str:
    return f"""
You are CareerOps Studio. Validate the CV against the Job Description and operating instructions.

Validation must be an internal ATS-style simulation only. Do not claim live website access.

Check:
- Manager CV vs Director CV selection
- Mandatory CV structure and order
- Exact mandatory role inclusion/exclusion
- Exact number of achievement lines per role
- Trade Compliance Manager EMEA must be exactly 8 achievements for Manager CV
- Head of Logistics EMEA must be exactly 7 achievements for Manager CV
- Every achievement line must be 170-190 characters including spaces
- No achievement may include ACHIEVEMENT:, Achievement:, Bullet:, Result:, Success:, KPI: or similar label
- No achievement may use bullet symbols or numbered lists
- No blank lines between achievements inside the same role
- Additional ATS Skills must be exactly 10 vertical lines, no bullets, no numbering, no blank lines
- Cover Letter length when present: 1,795-1,805 characters including spaces
- Cover Letter language must match the Job Description language unless explicitly requested otherwise
- Hard skills extracted from the Job Description
- Soft skills extracted from the Job Description
- Systems, tools, certifications and industry keywords extracted from the Job Description
- Whether Additional ATS Skills has exactly 10 skills aligned to the Job Description
- Whether the CV uses only Experience Repository facts
- Realistic simulated scores for LinkedIn-style, Jobscan-style, SkillSyncer-style, Resume Worded-style and Rezi-style checks
- Keyword Coverage Table exists
- Partial Match Keywords Table exists and contains 60-80% partial matches
- Missing Important Keywords Table exists
- Missing and partial keywords were considered during optimization
- Objective is zero deviations

Never claim 100% unless every critical keyword is naturally covered.
Return concise corrections only.
Language: {language}

{MANDATORY_CV_TYPE_RULES}

{MANDATORY_BULLET_COUNTS}

{ADDITIONAL_ATS_SKILLS_RULES}

{ATS_RULES}

{ATS_GAP_ANALYSIS_RULES}

CONTEXT FILES
{context}

JOB DESCRIPTION
{job_description}

CV TO VALIDATE
{cv_text}
"""

def show_header(provider_status: str, context_profile: str, selected_model: str):
    st.markdown(
        f"""
        <div class="header-container">
            <div class="header-top">
                <div>
                    <div class="header-brand">CareerOps Studio</div>
                    <div class="header-title">Executive CV Optimization</div>
                    <div class="header-subtitle">Map real achievements to job requirements, route Manager / Director CVs, validate ATS fit, and export client-ready files.</div>
                </div>
                <div class="status-badge">{provider_status}</div>
            </div>
            <div class="stats-grid">
                <div class="stat-card"><div class="stat-label">Primary Model</div><div class="stat-value">{selected_model}</div></div>
                <div class="stat-card"><div class="stat-label">Context Mode</div><div class="stat-value">{context_profile}</div></div>
                <div class="stat-card"><div class="stat-label">Routing</div><div class="stat-value">Manager / Director</div></div>
                <div class="stat-card"><div class="stat-label">Exports</div><div class="stat-value">PDF / DOCX / TXT</div></div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


st.set_page_config(page_title=APP_TITLE, page_icon="✨", layout="wide")
inject_css()

with st.sidebar:
    st.markdown("### ⚙️ Configuração")
    primary_model = st.selectbox("Modelo primário", MODEL_OPTIONS, index=0)
    failover_models = st.multiselect(
        "Modelos de fallback",
        MODEL_OPTIONS,
        default=["gemini-2.5-flash", "gemini-2.0-flash-lite", "gemini-1.5-flash"],
        help="Se o modelo primário falhar, CareerOps Studio tentará automaticamente o próximo modelo selecionado.",
    )
    custom_model = st.text_input("Modelo customizado (opcional)", value="")
    if custom_model.strip():
        failover_models = [custom_model.strip()] + failover_models
    context_profile = st.selectbox("Profundidade de contexto", ["Fast", "Balanced", "Deep"], index=1)
    st.caption("Múltiplas chaves de API suportadas.")
    st.divider()
    render_history_panel()
    st.divider()
    st.markdown("### 📂 Documentos de Origem")
    instructions_file = st.file_uploader("Instruções", type=["txt", "md", "docx", "pdf"], key="instructions")
    master_file = st.file_uploader("Repositório de Experiência", type=["txt", "md", "docx", "pdf"], key="master")
    manager_file = st.file_uploader("Template CV Manager", type=["txt", "md", "docx", "pdf"], key="manager")
    director_file = st.file_uploader("Template CV Director", type=["txt", "md", "docx", "pdf"], key="director")

instructions_text = read_uploaded_file(instructions_file) or load_first_available([
    "My_ideas/AI_CV_Instructions_Master_Rev7.docx", "AI_CV_Instructions_Master.docx", "AI_CV_Instructions_Master_Rev7.docx"
])
master_text = read_uploaded_file(master_file) or load_first_available([
    "Master Experience File.docx", "My_ideas/Master_Experience_File_Structured_Model.docx", "Master_Experience_File_Structured_Model.docx"
])
manager_text = read_uploaded_file(manager_file) or load_first_available([
    "cv manager x AI.docx", "My_ideas/Manager_CV_Template.docx", "Manager_CV_Template.docx"
])
director_text = read_uploaded_file(director_file) or load_first_available([
    "cv Director x AI.docx", "My_ideas/Director_CV_Template.docx", "Director_CV_Template.docx"
])
context = build_context(instructions_text, master_text, manager_text, director_text, context_profile)

provider_status = "Fallback ativado" if failover_models else "Modelo único"
show_header(provider_status, context_profile, primary_model)


if st.session_state.get("loaded_history_item"):
    loaded = st.session_state["loaded_history_item"]
    st.info(f"📂 Carregado do histórico: {loaded.get('created_at', '')} · {loaded.get('kind', '')} · {loaded.get('title', '')}")
    with st.expander("Resultado do histórico carregado"):
        st.markdown(loaded.get("result", ""))
        render_downloads(
            loaded.get("result", ""),
            f"careerops_history_{loaded.get('id', datetime.now().strftime('%Y%m%d_%H%M'))}",
        )

main_tab, validate_tab, files_tab = st.tabs(["💼 Workspace", "🔍 Validação ATS", "📋 Controle de Fonte"])

with main_tab:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">✍️ Criar CV Personalizado</div>', unsafe_allow_html=True)
    st.markdown('<p class="card-subtitle">Cole uma descrição de trabalho completa. O sistema garante formatação ATS compacta, skills ATS verticais, um achievement por linha, contagens de função exatas, carta de apresentação correspondente ao idioma e simulação ATS interna.</p>', unsafe_allow_html=True)
    job_description = st.text_area(
        "Descrição da vaga",
        height=360,
        placeholder="Cole a descrição completa da vaga aqui. Você pode incluir tabelas Markdown como: | Métrica | Pontuação |"
    )
    st.caption("Dica de entrada formatada: cole tabelas em formato Markdown. Exemplo: | Métrica | Pontuação |. Tabelas DOCX também são lidas de arquivos enviados.")
    c1, c2 = st.columns([1, 1])
    with c1:
        language = st.selectbox("Idioma de saída", ["Mesmo da descrição da vaga", "English", "Português do Brasil"])
    with c2:
        output_scope = st.selectbox(
            "Pacote de saída",
            [
                "ATS Validation Status, Tailored CV, Cover Letter (strict 1800 chars), Final ATS Report",
                "Tailored CV only with ATS Validation Status",
                "ATS Validation Status and Final ATS Report only",
            ],
        )
    generate = st.button("🚀 Gerar CV Pronto para Cliente", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if generate:
        if not job_description.strip():
            st.warning("Cole a descrição da vaga primeiro.")
        elif len(master_text.strip()) < 500:
            st.error("O Repositório de Experiência está ausente ou muito pequeno. Envie o arquivo Master Experience File real antes de gerar um CV.")
        else:
            prompt = strict_cv_prompt(context, job_description, language, output_scope)
            with st.spinner("Construindo CV e executando validação estrita..."):
                try:
                    result, used_model, attempts, strict_issues = generate_with_strict_validation(
                        prompt,
                        output_scope,
                        primary_model,
                        failover_models,
                    )
                    if strict_issues:
                        st.warning("A validação estrita ainda encontrou problemas após tentativas de reparo. Revise abaixo antes de enviar para o cliente.")
                        with st.expander("Problemas de validação estrita"):
                            for issue in strict_issues:
                                st.write("- " + issue)
                except Exception as exc:
                    st.error(str(exc))
                    st.stop()
            st.session_state["last_result"] = result
            st.session_state["last_job"] = job_description
            add_history_item(
                kind="CV",
                title=get_history_title(job_description, "CV Personalizado"),
                job_description=job_description,
                result=result,
                model=used_model,
                context_profile=context_profile,
            )
            st.markdown(f'<div class="audit-box">✅ Processado com: <b>{used_model}</b><br>Tentativas: {" → ".join(attempts)}</div>', unsafe_allow_html=True)
            st.markdown('<div class="result-container">', unsafe_allow_html=True)
            st.markdown(result)
            st.markdown('</div>', unsafe_allow_html=True)
            render_downloads(result, f"careerops_cv_{datetime.now().strftime('%Y%m%d_%H%M')}")

with validate_tab:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">🔍 Validar CV Existente</div>', unsafe_allow_html=True)
    st.markdown('<p class="card-subtitle">Use isto quando o cliente já tem um CV e quer verificar correspondência ATS, lacunas de palavras-chave e alinhamento de função.</p>', unsafe_allow_html=True)
    existing_cv = st.text_area("Texto do CV existente", height=260, value=st.session_state.get("last_result", ""))
    validation_job = st.text_area("Descrição da vaga para validação", height=240, value=st.session_state.get("last_job", ""))
    validate = st.button("🔎 Executar Validação ATS", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)
    if validate:
        if not existing_cv.strip() or not validation_job.strip():
            st.warning("Cole tanto o CV quanto a descrição da vaga.")
        else:
            prompt = ats_validation_prompt(context, existing_cv, validation_job, "Mesmo da descrição da vaga")
            with st.spinner("Validando ajuste ATS e testando fallback se necessário..."):
                try:
                    result, used_model, attempts = call_model(prompt, primary_model, failover_models)
                except Exception as exc:
                    st.error(str(exc))
                    st.stop()
            add_history_item(
                kind="Validation",
                title=get_history_title(validation_job, "Validação ATS"),
                job_description=validation_job,
                result=result,
                model=used_model,
                context_profile=context_profile,
            )
            st.markdown(f'<div class="audit-box">✅ Processado com: <b>{used_model}</b><br>Tentativas: {" → ".join(attempts)}</div>', unsafe_allow_html=True)
            st.markdown('<div class="result-container">', unsafe_allow_html=True)
            st.markdown(result)
            st.markdown('</div>', unsafe_allow_html=True)
            render_downloads(result, f"careerops_validation_{datetime.now().strftime('%Y%m%d_%H%M')}")

with files_tab:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📋 Arquivos de Origem Carregados</div>', unsafe_allow_html=True)
    st.markdown('<p class="card-subtitle">Confirme que o material de origem foi carregado corretamente. Tabelas Markdown e conteúdo de tabelas DOCX são exportados como células formatadas em DOCX e PDF.</p>', unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Instruções", f"{len(instructions_text):,}".replace(",", "."))
    c2.metric("Repositório de Experiência", f"{len(master_text):,}".replace(",", "."))
    c3.metric("Template Manager", f"{len(manager_text):,}".replace(",", "."))
    c4.metric("Template Director", f"{len(director_text):,}".replace(",", "."))
    st.markdown('</div>', unsafe_allow_html=True)
    with st.expander("👁️ Visualizar Instruções"):
        st.text(truncate_text(instructions_text, 6000))
    with st.expander("👁️ Visualizar Repositório de Experiência"):
        st.text(truncate_text(master_text, 6000))

st.markdown('<div class="footer">CareerOps Studio - Workspace privado de otimização de currículo executivo</div>', unsafe_allow_html=True)
